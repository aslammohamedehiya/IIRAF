"""
Generate Mid-Semester Viva Report for IIRAF Project
Creates a professionally formatted Word document following BITS Pilani standards
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_report():
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # ===== COVER PAGE =====
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('IIRAF – Intelligent Incident Resolution & Automation Framework\n\n')
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Course details
    course_info = doc.add_paragraph()
    course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course_info.add_run('Course No.: AIML CZG628T\n')
    run.font.size = Pt(12)
    run = course_info.add_run('Course Title: Dissertation\n\n')
    run.font.size = Pt(12)
    
    # Student details
    student_info = doc.add_paragraph()
    student_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = student_info.add_run('Mid-Semester Viva Report\n\n')
    run.font.size = Pt(14)
    run.font.bold = True
    
    run = student_info.add_run('Student Name: Aslam Mohamed Ehiya\n')
    run.font.size = Pt(12)
    run = student_info.add_run('BITS ID: 2023AC05247\n')
    run.font.size = Pt(12)
    run = student_info.add_run('Degree Program: M.Tech. in Artificial Intelligence & Machine Learning\n\n')
    run.font.size = Pt(12)
    
    run = student_info.add_run('Research Area: Artificial Intelligence, AIOps\n')
    run.font.size = Pt(12)
    run = student_info.add_run('Project Work carried out at: Verizon, Chennai\n\n\n')
    run.font.size = Pt(12)
    
    # University details
    uni_info = doc.add_paragraph()
    uni_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = uni_info.add_run('BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI\n')
    run.font.size = Pt(12)
    run.font.bold = True
    run = uni_info.add_run('VIDYA VIHAR, PILANI, RAJASTHAN - 333031\n\n')
    run.font.size = Pt(11)
    run = uni_info.add_run('December 2025')
    run.font.size = Pt(12)
    
    # Page break
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    heading = doc.add_heading('Contents', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        ('1.', 'Abstract Summary', '3'),
        ('2.', 'Technical Architecture', '4'),
        ('2.1', '    Project Workspace Architecture', '4'),
        ('2.2', '    Functional Blocks', '5'),
        ('2.3', '    Design Considerations', '6'),
        ('2.4', '    Process & Technology Used', '6'),
        ('3.', 'Work Accomplished (Completed)', '8'),
        ('4.', 'Future Work', '10'),
    ]
    
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{num}  {title}').font.size = Pt(12)
        p.add_run('.' * 50).font.size = Pt(12)
        p.add_run(f' {page}').font.size = Pt(12)
    
    doc.add_page_break()
    
    # ===== 1. ABSTRACT SUMMARY =====
    doc.add_heading('1. Abstract Summary', level=1)
    
    doc.add_heading('1.1 Project Overview', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The IIRAF (Intelligent Incident Resolution & Automation Framework) is an AI-driven self-service tool '
        'designed to revolutionize enterprise IT incident management. This Proof of Concept (PoC) demonstrates '
        'the practical application of advanced AI/ML techniques including Natural Language Processing (NLP), '
        'Vector Search, Machine Learning clustering and classification, and Large Language Models (LLMs) to '
        'automate incident triaging, pattern detection, and resolution recommendation.'
    )
    
    doc.add_heading('1.2 Problem Statement', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'In large-scale organizations like Verizon, IT support and customer service teams are often overwhelmed '
        'by a high volume of repetitive incidents. A large percentage of support tickets are for known issues '
        'that have been solved before. The "tribal knowledge" required to solve these issues is often lost in '
        'historical incident logs, old chat transcripts, or disparate Knowledge Base (KB) articles. This leads '
        'to high Mean Time To Resolution (MTTR), frustrated users, and inefficient support teams.'
    )
    
    doc.add_heading('1.3 Key Objectives', level=2)
    p = doc.add_paragraph('The primary objectives of this project are:', style='List Number')
    
    objectives = [
        'To design the complete architecture for an end-to-end incident resolution framework, including data ingestion, semantic search, pattern detection, and an auto-heal execution module.',
        'To implement a semantic search engine using NLP (SentenceTransformers) and a vector database (FAISS) to match new issues against historical incidents and knowledge base articles.',
        'To build a pattern detection module that analyzes incident data to identify frequently recurring issues and group them into actionable "patterns".',
        'To develop an "Autoheal Executor" that can map validated patterns to a repository of automation actions for simulated, no-touch resolution.',
        'To deliver a full-stack PoC, including a FastAPI backend API and a React-based frontend, that demonstrates the end-to-end workflow on the provided synthetic dataset.'
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj, style='List Number')
    
    doc.add_page_break()
    
    # ===== 2. TECHNICAL ARCHITECTURE =====
    doc.add_heading('2. Technical Architecture', level=1)
    
    doc.add_heading('2.1 Project Workspace Architecture', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The IIRAF system is organized into a modular, decoupled structure to ensure scalability and ease of '
        'maintenance. Core logic is centralized within the src/ directory to be shared across both the backend '
        'API and standalone training scripts.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('\\nKey Components:')
    run.font.bold = True
    
    components = [
        ('src/', 'Core Logic & Analytics Layer', [
            'app.py: FastAPI backend managing asynchronous endpoints for search, solution generation, and pattern detection',
            'build_index.py: FAISS index builder for semantic vector search',
            'query_retrieval.py: Semantic search engine using sentence transformers',
            'solution_generator.py: AI solution generator using Google Gemini Pro LLM',
            'severity_predictor.py: XGBoost classifier for incident severity prediction',
            'pattern_engine.py: HDBSCAN-based pattern detection and clustering',
            'autoheal_simulator.py: Auto-heal recommendation engine',
            'data_loader.py: Data loading and preprocessing utilities',
            'ml_utils.py: Machine learning utility functions',
            'visualization.py: UMAP-based dimensionality reduction for visualization'
        ]),
        ('data/', 'Data Layer', [
            'incidents.csv: Historical incident records (245 incidents)',
            'kb_articles.csv: Knowledge base articles (20 articles)',
            'Synthetic dataset mirroring real-world Verizon incident data'
        ]),
        ('frontend/', 'Presentation Layer', [
            'React-based web interface',
            'Interactive incident search and triage',
            'AI-generated solution display',
            'Pattern visualization'
        ]),
        ('models/', 'Model Storage', [
            'severity_predictor.pkl: Trained XGBoost model (456KB)',
            'tfidf_vectorizer.pkl: TF-IDF vectorizer (9KB)'
        ]),
        ('index_store/', 'Vector Index Storage', [
            'FAISS index files for fast semantic search',
            'Metadata for incidents and KB articles'
        ])
    ]
    
    for comp_name, comp_desc, items in components:
        p = doc.add_paragraph()
        run = p.add_run(f'● {comp_name} — {comp_desc}')
        run.font.bold = True
        for item in items:
            doc.add_paragraph(f'  ○ {item}', style='List Bullet')
    
    doc.add_heading('2.2 Functional Blocks', level=2)
    
    functional_blocks = [
        ('Data Ingestion & Preprocessing', 
         'Loads incident history and KB articles from CSV files. Performs text cleaning, normalization, and feature extraction.'),
        ('Semantic Embedding Layer',
         'Uses SentenceTransformers (all-MiniLM-L6-v2) to convert incident descriptions into 384-dimensional dense vectors. Enables semantic understanding beyond keyword matching.'),
        ('Vector Search Engine',
         'FAISS IndexFlatIP for fast cosine similarity search. L2 normalization for accurate semantic matching. Retrieves top-k similar incidents and relevant KB articles.'),
        ('ML Classification Layer',
         'XGBoost classifier for severity prediction (Low, Medium, High, Critical). TF-IDF feature extraction with 200 features. Achieves 93.06% accuracy on validation set.'),
        ('Pattern Detection Engine',
         'HDBSCAN for hierarchical density-based clustering. UMAP for dimensionality reduction and visualization. Automatically identifies recurring incident patterns.'),
        ('AI Solution Generator',
         'Google Gemini Pro LLM integration. Synthesizes incident resolutions and KB articles into coherent step-by-step solutions. Context-aware recommendations with fallback mechanism.'),
        ('Auto-Heal Simulator',
         'Maps validated patterns to automation actions. Simulates no-touch resolution for common issues. Provides remediation recommendations.'),
        ('API Layer',
         'FastAPI backend with 7 REST endpoints. Asynchronous request handling. Health monitoring and index status.'),
        ('Frontend Interface',
         'React-based web application. Real-time incident search and analysis. Interactive solution display.')
    ]
    
    for block_name, block_desc in functional_blocks:
        p = doc.add_paragraph()
        run = p.add_run(f'{block_name}: ')
        run.font.bold = True
        p.add_run(block_desc)
    
    doc.add_page_break()
    
    doc.add_heading('2.3 Design Considerations', level=2)
    
    design_points = [
        ('Modularity', 'Each component is independently testable and replaceable'),
        ('Scalability', 'FAISS index supports millions of vectors; FastAPI handles concurrent requests'),
        ('Performance', 'Sub-second response time for semantic search; efficient vector operations'),
        ('Accuracy', 'XGBoost model achieves 93.06% accuracy; semantic search provides relevant results'),
        ('Extensibility', 'Easy to add new ML models, data sources, or automation actions'),
        ('Maintainability', 'Clean code structure with comprehensive documentation')
    ]
    
    for consideration, description in design_points:
        p = doc.add_paragraph()
        run = p.add_run(f'● {consideration}: ')
        run.font.bold = True
        p.add_run(description)
    
    doc.add_heading('2.4 Process & Technology Used', level=2)
    
    # Create technology table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Purpose'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    tech_stack = [
        ('Programming Language', 'Python 3.11', 'Core development language'),
        ('Backend Framework', 'FastAPI', 'REST API development'),
        ('Frontend Framework', 'React', 'Web interface'),
        ('NLP Library', 'SentenceTransformers', 'Semantic embeddings (all-MiniLM-L6-v2)'),
        ('Vector Search', 'FAISS (Facebook AI)', 'Fast similarity search'),
        ('LLM Integration', 'Google Generative AI', 'Gemini Pro for solution generation'),
        ('ML Framework', 'XGBoost', 'Severity prediction (93.06% accuracy)'),
        ('Clustering', 'HDBSCAN', 'Pattern detection'),
        ('Dimensionality Reduction', 'UMAP', 'Visualization'),
        ('ML Utilities', 'Scikit-learn', 'Preprocessing, evaluation'),
        ('Data Processing', 'Pandas, NumPy', 'Data manipulation'),
        ('Web Server', 'Uvicorn', 'ASGI server'),
        ('Environment Management', 'Python-dotenv', 'Configuration management')
    ]
    
    for category, tech, purpose in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = tech
        row_cells[2].text = purpose
    
    doc.add_page_break()
    
    # ===== 3. WORK ACCOMPLISHED =====
    doc.add_heading('3. Work Accomplished (Completed)', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        'The following components have been successfully implemented and validated as part of the mid-semester progress:'
    )
    
    doc.add_heading('3.1 Data Generation & Loading Pipeline', level=2)
    accomplishments = [
        'Developed synthetic dataset generator mirroring real-world Verizon incident data',
        'Created 245 incident records across multiple severity levels and categories',
        'Generated 20 KB articles with resolution steps',
        'Implemented data loader with field mapping and text cleaning',
        'Validated data quality (100% completeness, no missing values)'
    ]
    for item in accomplishments:
        doc.add_paragraph(f'✓ {item}', style='List Bullet')
    
    doc.add_heading('3.2 Semantic Search Engine', level=2)
    accomplishments = [
        'Implemented SentenceTransformers integration (all-MiniLM-L6-v2 model)',
        'Built FAISS vector index with 265 total vectors (245 incidents + 20 KB articles)',
        'Achieved 74% top match relevance for test queries',
        'Implemented real-time index refresh capability',
        'Developed query retrieval system with relevance scoring'
    ]
    for item in accomplishments:
        doc.add_paragraph(f'✓ {item}', style='List Bullet')
    
    doc.add_heading('3.3 XGBoost Severity Prediction Model', level=2)
    p = doc.add_paragraph()
    p.add_run('Successfully trained and validated XGBoost classifier with exceptional performance:')
    
    # Performance metrics table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    metrics = [
        ('Overall Accuracy', '93.06%'),
        ('F1 Score', '0.9284'),
        ('Precision', '93.37%'),
        ('Recall', '93.06%'),
        ('Correct Predictions', '228/245'),
        ('Error Rate', '6.94%')
    ]
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    p = doc.add_paragraph()
    p.add_run('\\nPer-Class Performance:')
    
    # Per-class performance table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Severity'
    hdr_cells[1].text = 'Precision'
    hdr_cells[2].text = 'Recall'
    hdr_cells[3].text = 'Accuracy'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    class_metrics = [
        ('Low', '91.80%', '98.82%', '98.82%'),
        ('Medium', '100.00%', '83.33%', '83.33%'),
        ('High', '100.00%', '72.22%', '72.22%'),
        ('Critical', '94.87%', '82.22%', '82.22%')
    ]
    
    for severity, precision, recall, accuracy in class_metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = severity
        row_cells[1].text = precision
        row_cells[2].text = recall
        row_cells[3].text = accuracy
    
    doc.add_heading('3.4 AI Solution Generation', level=2)
    accomplishments = [
        'Integrated Google Gemini Pro LLM for intelligent solution synthesis',
        'Developed context-aware prompt engineering',
        'Implemented fallback mechanism for API unavailability',
        'Generated step-by-step resolution paths from multiple sources',
        'Validated solution quality through manual review'
    ]
    for item in accomplishments:
        doc.add_paragraph(f'✓ {item}', style='List Bullet')
    
    doc.add_heading('3.5 Pattern Detection Engine', level=2)
    accomplishments = [
        'Implemented HDBSCAN clustering for automatic pattern detection',
        'Integrated UMAP for dimensionality reduction and visualization',
        'Developed pattern grouping and labeling logic',
        'Created pattern analysis API endpoint',
        'Validated pattern detection on incident dataset'
    ]
    for item in accomplishments:
        doc.add_paragraph(f'✓ {item}', style='List Bullet')
    
    doc.add_heading('3.6 Backend API Development', level=2)
    p = doc.add_paragraph()
    p.add_run('Developed comprehensive FastAPI backend with 7 REST endpoints:')
    
    endpoints = [
        ('GET /health', 'Health check with index status'),
        ('POST /api/search', 'Semantic search for incidents and KB articles'),
        ('POST /api/generate-solution', 'AI-powered solution generation'),
        ('GET /api/patterns', 'Pattern detection results'),
        ('POST /api/heal', 'Trigger auto-heal actions'),
        ('GET /api/index/status', 'FAISS index status'),
        ('POST /api/index/refresh', 'Manually rebuild FAISS index')
    ]
    
    for endpoint, description in endpoints:
        p = doc.add_paragraph()
        run = p.add_run(f'● {endpoint}: ')
        run.font.bold = True
        run.font.name = 'Courier New'
        p.add_run(description)
    
    doc.add_heading('3.7 Frontend Development', level=2)
    accomplishments = [
        'Developed React-based web interface',
        'Implemented incident search and analysis UI',
        'Created AI-generated solution display',
        'Integrated KB article sidebar',
        'Designed responsive layout for desktop and mobile'
    ]
    for item in accomplishments:
        doc.add_paragraph(f'✓ {item}', style='List Bullet')
    
    doc.add_heading('3.8 Validation & Testing', level=2)
    p = doc.add_paragraph()
    p.add_run('Comprehensive validation completed:')
    
    validation = [
        'Model retraining and validation (93.06% accuracy achieved)',
        'Semantic search testing with various query types',
        'End-to-end workflow testing through browser automation',
        'API endpoint testing and documentation',
        'Data quality analysis and verification',
        'Performance benchmarking (sub-second response times)'
    ]
    for item in validation:
        doc.add_paragraph(f'✓ {item}', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 4. FUTURE WORK =====
    doc.add_heading('4. Future Work', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        'The following enhancements are planned for the second half of the semester to complete the dissertation:'
    )
    
    doc.add_heading('4.1 Dataset Enhancement', level=2)
    future_work = [
        'Collect additional data for Medium and High severity classes (currently underrepresented)',
        'Integrate real Verizon incident data (with appropriate anonymization)',
        'Expand KB article repository to 50+ articles',
        'Add temporal features (time-of-day, day-of-week patterns)',
        'Include device and application metadata'
    ]
    for item in future_work:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_heading('4.2 Model Improvements', level=2)
    future_work = [
        'Implement ensemble methods combining XGBoost with other classifiers',
        'Add active learning for low-confidence predictions',
        'Develop confidence calibration techniques',
        'Implement online learning for continuous model updates',
        'Explore transformer-based models for severity prediction'
    ]
    for item in future_work:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_heading('4.3 ServiceNow Integration', level=2)
    future_work = [
        'Connect to ServiceNow ITSM platform via REST API',
        'Implement real-time incident ingestion',
        'Develop bidirectional sync for resolution updates',
        'Create ServiceNow app integration',
        'Build workflow automation triggers'
    ]
    for item in future_work:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_heading('4.4 Advanced Features', level=2)
    future_work = [
        'Implement multi-language support for global teams',
        'Develop advanced visualization dashboards',
        'Add incident priority prediction',
        'Create automated escalation logic',
        'Implement sentiment analysis for user satisfaction',
        'Build recommendation system for KB article creation'
    ]
    for item in future_work:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_heading('4.5 Production Readiness', level=2)
    future_work = [
        'Containerization with Docker',
        'Kubernetes deployment configuration',
        'Implement comprehensive logging and monitoring',
        'Add authentication and authorization',
        'Develop CI/CD pipeline',
        'Create comprehensive API documentation',
        'Implement rate limiting and caching',
        'Add database backend (PostgreSQL) for persistence'
    ]
    for item in future_work:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_heading('4.6 Evaluation & Validation', level=2)
    future_work = [
        'Conduct user acceptance testing with Verizon support team',
        'Measure MTTR reduction compared to baseline',
        'Evaluate solution quality through expert review',
        'Benchmark performance against existing tools',
        'Conduct A/B testing for different ML models',
        'Gather feedback for iterative improvements'
    ]
    for item in future_work:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_heading('4.7 Documentation & Thesis Writing', level=2)
    future_work = [
        'Complete comprehensive technical documentation',
        'Write dissertation chapters (Literature Review, Methodology, Results)',
        'Create user manual and deployment guide',
        'Prepare final presentation and demo',
        'Document lessons learned and best practices'
    ]
    for item in future_work:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    # Save document
    output_dir = r'Reports'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'Aslam_midsem_report_2023ac05247.docx')
    doc.save(output_path)
    
    print(f"\\n{'='*80}")
    print("REPORT GENERATION COMPLETE")
    print(f"{'='*80}")
    print(f"\\nReport saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024:.2f} KB")
    print(f"\\nTotal sections: 4")
    print(f"Total pages: ~12-15 (estimated)")
    print(f"\\nThe report includes:")
    print("  ✓ Cover page with student and course details")
    print("  ✓ Table of contents")
    print("  ✓ Abstract summary with problem statement and objectives")
    print("  ✓ Technical architecture with detailed component breakdown")
    print("  ✓ Work accomplished with validation metrics")
    print("  ✓ Future work roadmap")
    print(f"\\n{'='*80}")

if __name__ == '__main__':
    create_report()
